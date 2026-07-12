import docx

def replace_text_in_docx(doc_path, save_path, replacements):
    doc = docx.Document(doc_path)
    
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, val)
                    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in replacements.items():
                        if key in p.text:
                            p.text = p.text.replace(key, val)

    doc.save(save_path)
    print("Template saved to", save_path)

if __name__ == "__main__":
    replacements = {
        "Name :": "Name : {~applicantName~}",
        "Code No :": "Code No : {~employeeCode~}",
        "From :": "From : {~fromDate~}",
        "To   :": "To : {~toDate~}",
        "Designation :": "Designation : {~designationGrade~}",
        "No. of days required :": "No. of days required : {~numberOfDays~}",
        "Division  / Section :": "Division / Section : {~divisionSection~}",
        "Prefixing : -": "Prefixing : - {~prefixing~}",
        "Nature of leave  :": "Nature of leave  : {~natureOfLeave~}",
        "Suffixing : -": "Suffixing : - {~suffixing~}",
        "Purpose:": "Purpose: {~purpose~}",
        "Address during leave :": "Address during leave : {~addressDuringLeave~}",
        "Date:": "Date: {~applicationDate~}",
        "Intervening closed holidays : -": "Intervening closed holidays : - {~interveningClosedHolidays~}",
        "Casual leave taken so far :": "Casual leave taken so far : {~clTakenSoFar~}",
    }
    
    in_file = r"C:\Users\vishn\OneDrive\Desktop\ISRO\Regional-Remote\nrsc-slms\form\Casual Leave (1).docx"
    out_file = r"C:\Users\vishn\OneDrive\Desktop\ISRO\Regional-Remote\nrsc-slms\backend\src\docx_templates\casual_leave_rrsc.docx"
    
    replace_text_in_docx(in_file, out_file, replacements)
